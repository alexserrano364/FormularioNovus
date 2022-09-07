import os.path
from datetime import datetime
from tkinter import *
from tkinter import font
from tkinter import messagebox

import openpyxl
from PIL import ImageTk, Image
from openpyxl import *
from openpyxl.styles import Font, PatternFill

from docx import Document
from docx.enum.section import WD_ORIENTATION

meses = {1: "Enero", 2: "Febrero", 3: "Marzo", 4: "Abril", 5: "Mayo", 6: "Junio", 7: "Julio", 8: "Agosto",
         9: "Septiembre", 10: "Octubre", 11: "Noviembre", 12: "Diciembre"}

blue = "#6bd8f0"
fontName = "Montserrat"


def obtenerHora():
    today = datetime.now()
    hour = today.hour
    minute = today.minute
    return [hour, minute]


def obtenerFecha():
    today = datetime.now()
    year = today.year
    month = today.month
    day = today.day
    return [day, month, year]


def prepararEncabezado(sheet):
    # Toooodas las celdas
    sheet['A3'] = "ELABORO"

    sheet['A6'] = "INGRESOS"
    sheet['A7'] = "FOLIO"
    sheet['B7'] = "PACIENTE"
    sheet['C7'] = "MOTIVO"
    sheet['D7'] = "EFECTIVO"
    sheet['E7'] = "VAUCHER"

    sheet['E2'] = "TOTAL EFECTIVO"
    sheet['F2'] = "=SUM(D:D)"
    sheet['E3'] = "TOTAL VAUCHER"
    sheet['F3'] = "=SUM(E:E)"
    sheet['E4'] = "TOTAL"
    sheet['F4'] = "=F2+F3"

    sheet['H3'] = "SALDO"
    sheet['I3'] = "=F4-M3"

    sheet['J6'] = "EGRESOS"
    sheet['J7'] = "# VALE"
    sheet['K7'] = "MOTIVO"
    sheet['L7'] = "IMPORTE"

    sheet['L3'] = "TOTAL PAGOS"
    sheet['M3'] = "=SUM(L:L)"

    # Hace a las celdas grises con texto blanco
    titulos = ["A3", "A6", "A7", "B7", "C7", "D7", "E7", "E2", "E3", "E4", "H3", "J6", "J7", "K7", "L7", "L3"]
    greyFill = PatternFill(start_color='404040', end_color='404040', fill_type='solid')
    for cell in titulos:
        activeCell = sheet[cell]
        activeCell.font = Font(bold=True, color="FFFFFF")
        activeCell.fill = greyFill

    # Une celdas
    sheet.merge_cells("A6:B6")
    sheet.merge_cells("J6:K6")

    # Ancho de columnas
    sheet.column_dimensions['A'].width = 20.5
    sheet.column_dimensions['B'].width = 32.5
    sheet.column_dimensions['C'].width = 20.5
    sheet.column_dimensions['D'].width = 14.17
    sheet.column_dimensions['E'].width = 14.17
    sheet.column_dimensions['K'].width = 32.5
    sheet.column_dimensions['L'].width = 14.17

    # TODO: Formato de número


def conseguirDatos(botones):
    datos = []
    for boton in botones:           # Recorre todos los botones y los convierte en texto
        dato = boton.get()
        datos.append(dato)
    return datos


def limpiarDatos(datos):
    datos[0] = datos[0].upper()  # Pasa el nombre a mayúsculas

    datos[1] = datos[1].replace(" ", "")  # Se deshace de espacios y guiones en el núm. tel.
    datos[1] = datos[1].replace("-", "")

    datos[4] = datos[4].upper()  # Pasa el servicio a mayúsculas


def datosSonValidos(datos):

    try:                                # Checa que el número telefónico sea un número
        int(datos[1])
    except ValueError:
        messagebox.showinfo("Alerta", "El número telefónico debe ser un número")
        return False

    try:                                # Checa que el código postal sea un número
        int(datos[3])
    except ValueError:
        messagebox.showinfo("Alerta", "El código postal debe ser un número")
        return False

    try:                                # Checa que el importe sea un número
        float(datos[5])
    except ValueError:
        messagebox.showinfo("Alerta", "El importe debe ser un número")
        return False
    else:
        datos[5] = "$" + datos[5]

    if datos[6] == "Seleccionar":       # Checa que el turno haya sido seleccionado
        messagebox.showinfo("Alerta", "Seleccionar Turno")
        return False

    if datos[7] == "Seleccionar":       # Checa que el tipo de pago haya sido seleccionado
        messagebox.showinfo("Alerta", "Seleccionar Tipo de Pago")
        return False

    return True


def crearEImprimirDocumento(datos):
    document = Document()

    pic = "./img/NovusLogo.jpeg"
    style = document.styles['Normal']
    font = style.font
    font.name = 'Consolas'

    section = document.sections[-1]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height
    table = document.add_table(rows=1, cols=3)

    for row in table.rows:
        for cell in row.cells:
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(pic, height=1000000)
            paragraph = cell.add_paragraph(("\n\nFOLIO: %s\n\n\n\n" +
                                            "NOMBRE: %s\n\n\n\n" +
                                            "SERVICIO: %s\n\n\n\n" +
                                            "IMPORTE: %s\n\n\n\n" +
                                            "FECHA: %s\n\n" +
                                            "HORA: %s\n\n") % (datos[7], datos[0], datos[4], datos[5], datos[8], datos[6]))
            paragraph.style = document.styles['Normal']
    document.save("Ticket.docx")
    os.startfile("Ticket.docx", "print")


def crearExcel(root, botones):
    datos = conseguirDatos(botones)
    limpiarDatos(datos)
    if datosSonValidos(datos):                # Solo funciona si los datos están bien

        # Se obtienen datos de fecha y hora
        fecha = obtenerFecha()
        hora = obtenerHora()
        if 0 <= hora[0] < 7:  # Si es turno de madrugada, es parte del corte del día anterior
            fecha[0] -= 1

        # Se crea el Excel del turno, si ya existe solo se carga
        turno = datos[6][:1]
        title = './Cortes de Caja/Corte Caja %d de %s del %d - %s.xlsx' % (fecha[0], meses[fecha[1]], fecha[2], turno)
        if os.path.exists(title):
            workbook = load_workbook(title)
            sheet = workbook.active
        else:
            workbook = openpyxl.Workbook()
            sheet = workbook.active
            prepararEncabezado(sheet)

        # Se crea el folio y se da formato a la fecha y hora
        seriado = len(sheet['H'])-6
        folio = "%03d-%04d%02d%02d%s" % (seriado, fecha[2], fecha[1], fecha[0], turno)
        stringHora = "%02d:%02d" % (hora[0], hora[1])
        stringFecha = "%02d/%02d/%02d" % (fecha[0], fecha[1], fecha[2])

        # Se hace el formato para el corte de caja
        datosCorteDeCaja = [folio, datos[0], datos[4]]

        if datos[7] == "Tarjeta":               # Revisa cómo fue el pago
            datosCorteDeCaja.append("")
            datosCorteDeCaja.append(datos[5])
        else:
            datosCorteDeCaja.append(datos[5])
            datosCorteDeCaja.append("")
        datos.remove(datos[7])

        sheet.append(datosCorteDeCaja)
        workbook.save(title)
        workbook.close()

        # Se añade la info a la base general que ya existe
        title = "Base General.xlsx"
        workbook = load_workbook(title)
        sheet = workbook.active
        datos.remove(datos[6])
        datos.append(stringHora)
        datos.append(folio)
        datos.append(stringFecha)
        sheet.append(datos)
        workbook.save(title)
        workbook.close()

        messagebox.showinfo("Éxito", "Datos agregados")
        crearEImprimirDocumento(datos)
        root.quit()


def crearInfoVentana(root):

    padx = 35
    pady = 20

    fuenteTitulo = font.Font(bold=True, font=(fontName, 36))
    fuenteSubtitulo = font.Font(italic=True, font=(fontName, 24))
    fuenteLabels = font.Font(font=(fontName, 14))
    fuenteMenu = font.Font(font=(fontName, 10))
    fuenteBoton = font.Font(font=(fontName, 12))

    labelBienvenido = Label(root, text="¡Bienvenido!", font=fuenteTitulo)
    labelSubtitulo = Label(root, text="Sistema de Corte de Caja Novus Médica", font=fuenteSubtitulo)
    fecha = obtenerFecha()
    hora = obtenerHora()
    labelFechaHora = Label(root, text="La fecha de hoy es %d de %s. " % (fecha[0], meses[fecha[1]])
                           + "La hora es %02d:%02d. (Versión 1.0)" % (hora[0], hora[1]), font=fuenteMenu)

    labelNombre = Label(root, text="Nombre:", font=fuenteLabels)
    nombre = Entry(root, width=85)

    labelTurno = Label(root, text="Turno:", font=fuenteLabels)
    turno = StringVar()
    turno.set("Seleccionar")
    turnoMenu = OptionMenu(root, turno, "Matutino", "Vespertino", "Nocturno")
    turnoMenu.config(font=fuenteBoton)
    turnoDropMenu = root.nametowidget(turnoMenu.menuname)
    turnoDropMenu.config(font=fuenteMenu)                 # Se le empata la fuente de las labels a las opciones

    labelServicio = Label(root, text="Servicio:", font=fuenteLabels)
    servicio = Entry(root, width=85)

    labelImporte = Label(root, text="Importe:", font=fuenteLabels)
    importe = Entry(root, width=15)

    labelPago = Label(root, text="Método de pago:", font=fuenteLabels)
    tipoPago = StringVar()
    tipoPago.set("Seleccionar")
    pagoMenu = OptionMenu(root, tipoPago, "Efectivo", "Tarjeta")
    pagoMenu.config(font=fuenteBoton)
    turnoDropPago = root.nametowidget(pagoMenu.menuname)
    turnoDropPago.config(font=fuenteMenu)

    labelCelular = Label(root, text="Número Celular:", font=fuenteLabels)
    celular = Entry(root, width=85)

    labelCorreo = Label(root, text="Correo Electrónico:", font=fuenteLabels)
    correo = Entry(root, width=85)

    labelCP = Label(root, text="Código Postal:", font=fuenteLabels)
    CP = Entry(root, width=15)

    imgLogo = ImageTk.PhotoImage(Image.open("./img/NovusLogo.jpeg"))
    labelImagen = Label(image=imgLogo)

    botones = [nombre, celular, correo, CP, servicio, importe, turno, tipoPago]

    boton = Button(root, text="Insertar", padx=10,
                   command=lambda: crearExcel(root, botones),
                   bg=blue, fg="black", font=fuenteBoton, width=20, height=4)

    # Insertar labels
    labelBienvenido.grid(row=0, column=1, padx=padx, pady=pady, sticky="W")
    labelSubtitulo.grid(row=1, column=1, padx=padx, pady=10, sticky="W")
    labelFechaHora.grid(row=12, column=1, padx=padx, pady=pady, sticky="SW")
    labelNombre.grid(row=2, column=1, padx=padx, pady=pady, sticky="W")
    nombre.grid(row=3, column=1, padx=padx, sticky="W")
    labelTurno.grid(row=2, column=3, padx=padx, pady=pady, sticky="W")
    turnoMenu.grid(row=3, column=3, padx=padx, sticky="W")
    labelServicio.grid(row=5, column=1, padx=padx, pady=pady, sticky="W")
    servicio.grid(row=6, column=1, padx=padx, sticky="W")
    labelImporte.grid(row=5, column=3, padx=padx, pady=pady, sticky="W")
    importe.grid(row=6, column=3, padx=padx, sticky="W")
    labelPago.grid(row=8, column=3,padx=padx, sticky="W")
    pagoMenu.grid(row=9, column=3,padx=padx, sticky="W")
    labelCelular.grid(row=8, column=1, padx=padx, pady=pady, sticky="W")
    celular.grid(row=9, column=1, padx=padx, sticky="W")
    labelCP.grid(row=10, column=3, padx=padx, pady=pady, sticky="W")
    CP.grid(row=11, column=3, padx=padx, sticky="W")
    labelCorreo.grid(row=10, column=1, padx=padx, pady=pady, sticky="W")
    correo.grid(row=11, column=1, padx=padx, sticky="W")
    boton.grid(row=12, column=3, padx=padx, pady=45)
    # labelImagen.grid(row=12, column=1)


def main():
    # Se crea la ventana
    root = Tk()
    root.title("Sistema de Corte de Caja Novus Médica")
    root.iconbitmap('./img/caduceus.ico')
    crearInfoVentana(root)
    root.mainloop()


main()
